"""
DOCX Extractor - Sprint 12.2

Extracts text from Word (.docx) documents for document processing.
"""

from pathlib import Path
from typing import Dict, Any
import logging

logger = logging.getLogger(__name__)


class DOCXExtractor:
    """Extract text from DOCX files."""

    def extract_text(self, file_path: Path) -> Dict[str, Any]:
        """
        Extract text from a DOCX file.

        Args:
            file_path: Path to the DOCX file

        Returns:
            {
                "success": bool,
                "text": str,
                "paragraph_count": int,
                "table_count": int,
                "error": str | None
            }
        """
        result = {
            "success": False,
            "text": "",
            "paragraph_count": 0,
            "table_count": 0,
            "error": None,
        }

        try:
            from docx import Document

            doc = Document(file_path)
            text_parts = []

            # Paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            result["paragraph_count"] = len(doc.paragraphs)

            # Tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))
                result["table_count"] += 1

            result["text"] = "\n\n".join(text_parts)
            result["success"] = True

        except ImportError:
            result["error"] = "python-docx not installed. Run: pip install python-docx"
            logger.error(result["error"])
        except Exception as e:
            result["error"] = str(e)
            logger.error(f"DOCX extraction failed for {file_path}: {e}")

        return result
